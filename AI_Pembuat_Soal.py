import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import markdown
from htmldocx import HtmlToDocx
import re
import requests
import urllib.parse
import time
import json

# =====================================================
# 1. KONFIGURASI HALAMAN
# =====================================================
st.set_page_config(
    page_title="SmartQuiz AI",
    page_icon="⚡",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Plus+Jakarta+Sans:ital,wght@0,400;0,500;0,600;0,700;0,800;1,400&display=swap');
html, body, [class*="css"] { font-family: 'Plus Jakarta Sans', sans-serif !important; }

[data-testid="stAppViewContainer"] {
    background: linear-gradient(145deg, #EEF2FF 0%, #F0FAFA 50%, #FAF5FF 100%);
    min-height: 100vh;
}
[data-testid="stSidebar"] {
    background: linear-gradient(180deg, #1E1B4B 0%, #312E81 100%);
    border-right: none;
}
[data-testid="stSidebar"] * { color: #E0E7FF !important; }
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3 { color: #FFFFFF !important; }
[data-testid="stSidebar"] .stTextInput input {
    background: rgba(255,255,255,0.1) !important;
    border: 1px solid rgba(255,255,255,0.2) !important;
    color: white !important;
    border-radius: 10px !important;
}
[data-testid="stSidebar"] .stTextInput input::placeholder { color: rgba(255,255,255,0.5) !important; }
[data-testid="stSidebar"] hr { border-color: rgba(255,255,255,0.15) !important; }
[data-testid="stSidebar"] .stSelectbox label { color: #A5B4FC !important; font-size: 0.82rem !important; }
[data-testid="stSidebar"] [data-baseweb="select"] > div {
    background: rgba(255,255,255,0.1) !important;
    border: 1px solid rgba(255,255,255,0.2) !important;
    border-radius: 10px !important;
}

[data-testid="block-container"] {
    background: rgba(255,255,255,0.80);
    backdrop-filter: blur(24px);
    -webkit-backdrop-filter: blur(24px);
    border-radius: 28px;
    padding: 2.5rem 3.5rem;
    box-shadow: 0 25px 50px rgba(79,70,229,0.07), 0 8px 16px rgba(0,0,0,0.04);
    border: 1px solid rgba(255,255,255,0.7);
    margin-top: 1.5rem; margin-bottom: 2rem;
}
.title-text {
    background: linear-gradient(135deg, #4338CA 0%, #7C3AED 50%, #EC4899 100%);
    -webkit-background-clip: text; -webkit-text-fill-color: transparent;
    font-weight: 800; font-size: 3.2rem; letter-spacing: -1.5px; line-height: 1.15;
}
.subtitle-badge {
    display: inline-block;
    background: linear-gradient(135deg, #4F46E5, #7C3AED);
    color: white !important; font-size: 0.72rem; font-weight: 700;
    letter-spacing: 0.12em; text-transform: uppercase;
    padding: 4px 14px; border-radius: 100px; margin-bottom: 0.75rem;
}
.section-header {
    font-size: 1.05rem; font-weight: 700; color: #1E1B4B;
    border-left: 4px solid #4F46E5; padding: 0.4rem 0 0.4rem 0.8rem;
    margin-bottom: 1rem; margin-top: 0.5rem;
}
.provider-badge {
    display: inline-block; font-size: 0.7rem; font-weight: 700;
    padding: 2px 10px; border-radius: 100px; margin-left: 6px; vertical-align: middle;
}
.badge-gemini  { background: #EEF2FF; color: #4F46E5; border: 1px solid #C7D2FE; }
.badge-groq    { background: #FFF7ED; color: #C2410C; border: 1px solid #FED7AA; }
.badge-openrouter { background: #F0FDF4; color: #166534; border: 1px solid #BBF7D0; }

div[data-baseweb="input"] > div,
div[data-baseweb="select"] > div,
div[data-baseweb="textarea"] > div {
    border-radius: 12px !important; background-color: #FAFBFF !important;
    border: 1.5px solid #DDE3F5 !important; transition: all 0.25s ease;
}
div[data-baseweb="input"] > div:focus-within,
div[data-baseweb="select"] > div:focus-within,
div[data-baseweb="textarea"] > div:focus-within {
    border-color: #4F46E5 !important;
    box-shadow: 0 0 0 3px rgba(79,70,229,0.15) !important;
    background-color: #FFFFFF !important;
}
div[data-testid="stNumberInput"] input {
    border-radius: 10px !important; text-align: center; font-weight: 700; font-size: 1.1rem;
}
.stButton > button {
    background: linear-gradient(135deg, #4F46E5 0%, #7C3AED 100%);
    color: white !important; font-weight: 700; font-size: 1.05rem;
    border-radius: 14px; padding: 0.85rem 1.5rem; width: 100%; border: none;
    box-shadow: 0 6px 20px rgba(79,70,229,0.35);
    transition: all 0.3s cubic-bezier(0.4,0,0.2,1); letter-spacing: 0.01em;
}
.stButton > button:hover {
    transform: translateY(-3px) scale(1.01);
    box-shadow: 0 12px 30px rgba(79,70,229,0.45); color: white !important;
}
.stDownloadButton > button {
    background: linear-gradient(135deg, #059669 0%, #0D9488 100%) !important;
    color: white !important; font-weight: 700; border-radius: 12px; border: none !important;
    box-shadow: 0 4px 15px rgba(5,150,105,0.3) !important; transition: all 0.3s ease !important;
}
.stDownloadButton > button:hover {
    transform: translateY(-2px) !important; box-shadow: 0 8px 25px rgba(5,150,105,0.4) !important;
}
[data-baseweb="tab-list"] { gap: 0.5rem; border-bottom: 2px solid #E8EDFF; margin-bottom: 1.5rem; }
[data-baseweb="tab"] {
    font-weight: 600; font-size: 0.95rem; color: #6B7280;
    padding: 0.6rem 1rem; border-radius: 8px 8px 0 0; transition: all 0.2s ease;
}
[data-baseweb="tab"]:hover { color: #4F46E5; background: rgba(79,70,229,0.05); }
[data-baseweb="tab"][aria-selected="true"] { color: #4F46E5; }
[data-baseweb="tab-highlight"] {
    background: linear-gradient(90deg, #4F46E5, #7C3AED); height: 3px; border-radius: 3px 3px 0 0;
}
.stAlert { border-radius: 12px !important; }
hr { border-color: #E8EDFF !important; margin: 1.5rem 0 !important; }
.stat-card {
    background: linear-gradient(135deg, #EEF2FF 0%, #F5F3FF 100%);
    border: 1px solid #C7D2FE; border-radius: 14px; padding: 1rem 1.25rem; text-align: center;
}
.stat-number { font-size: 2rem; font-weight: 800; color: #4F46E5; line-height: 1; }
.stat-label  { font-size: 0.8rem; color: #6366F1; font-weight: 600; margin-top: 4px; }
[data-testid="stExpander"] { border: 1px solid #DDE3F5 !important; border-radius: 14px !important; overflow: hidden; }

@media (max-width: 768px) {
    [data-testid="block-container"] { padding: 1.2rem 1rem; border-radius: 16px; margin-top: 0.25rem; }
    .title-text { font-size: 2rem; }
}
</style>
""", unsafe_allow_html=True)


# =====================================================
# 2. HEADER
# =====================================================
st.markdown("""
<div style="text-align:center; margin-bottom:2.5rem;">
    <span class="subtitle-badge">Multi-Provider · 100% Free · AI-Powered</span>
    <h1 class="title-text">SmartQuiz AI ⚡</h1>
    <p style="color:#64748B; font-size:1.1rem; max-width:680px; margin:0.6rem auto 0; line-height:1.7;">
        Generator soal evaluasi otomatis — pilih provider AI favoritmu: 
        <strong>Gemini</strong>, <strong>Groq</strong>, atau <strong>OpenRouter</strong>.
        Lengkap dengan kunci jawaban, kisi-kisi, kartu soal, dan ekspor Word.
    </p>
</div>
""", unsafe_allow_html=True)


# =====================================================
# 3. FUNGSI EXPORT KE WORD
# =====================================================
def export_to_docx(judul, info, hasil_ai):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    title_para = doc.add_heading(judul, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.runs[0]
    run.font.color.rgb = RGBColor(0x2D, 0x31, 0x9E)
    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate([
        ("Mata Pelajaran", info["mapel"]),
        ("Kelas", info["kelas"]),
        ("Topik / Tujuan Pembelajaran", info["topik"]),
    ]):
        row = table.rows[i]
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    doc.add_paragraph()

    parts = re.split(r'!\[.*?\]\((.*?)\)', hasil_ai)
    new_parser = HtmlToDocx()
    for i, part in enumerate(parts):
        if i % 2 == 0:
            if part.strip():
                clean_text = part.replace("```markdown", "").replace("```", "")
                html_text = markdown.markdown(clean_text, extensions=['tables', 'nl2br', 'sane_lists'])
                new_parser.add_html_to_document(html_text, doc)
        else:
            url = part.strip()
            try:
                resp = requests.get(url, timeout=20, headers={'User-Agent': 'Mozilla/5.0'}, allow_redirects=True)
                if resp.status_code == 200:
                    doc.add_picture(BytesIO(resp.content), width=Inches(4.0))
                else:
                    doc.add_paragraph(f"[Gambar tidak tersedia – kode error: {resp.status_code}]")
            except Exception:
                doc.add_paragraph("[Gambar tidak dapat dimuat.]")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =====================================================
# 4. FUNGSI GENERATE — MULTI PROVIDER
# =====================================================

# ── Gemini ──────────────────────────────────────────
GEMINI_MODELS = [
    "gemini-2.0-flash", "gemini-2.0-flash-lite",
    "gemini-1.5-flash-latest", "gemini-1.5-flash", "gemini-1.5-pro",
]

def get_best_gemini_model(api_key: str) -> str:
    try:
        genai.configure(api_key=api_key)
        available = {m.name.split("/")[-1] for m in genai.list_models()
                     if "generateContent" in m.supported_generation_methods}
        for m in GEMINI_MODELS:
            if m in available:
                return m
        for name in available:
            if "flash" in name:
                return name
    except Exception:
        pass
    return "gemini-1.5-flash"

def generate_gemini(prompt: str, api_key: str) -> tuple[str, str]:
    model_name = get_best_gemini_model(api_key)
    genai.configure(api_key=api_key)
    gen_cfg = genai.GenerationConfig(temperature=0.85, top_p=0.95, max_output_tokens=8192)
    safety = [
        {"category": "HARM_CATEGORY_HARASSMENT",        "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_HATE_SPEECH",        "threshold": "BLOCK_NONE"},
        {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",  "threshold": "BLOCK_ONLY_HIGH"},
        {"category": "HARM_CATEGORY_DANGEROUS_CONTENT",  "threshold": "BLOCK_NONE"},
    ]
    model = genai.GenerativeModel(model_name, generation_config=gen_cfg, safety_settings=safety)
    return model.generate_content(prompt).text, model_name

# ── Groq ─────────────────────────────────────────────
GROQ_MODELS = [
    "llama-3.3-70b-versatile",
    "llama-3.1-70b-versatile",
    "mixtral-8x7b-32768",
    "gemma2-9b-it",
]

def generate_groq(prompt: str, api_key: str, model_choice: str) -> tuple[str, str]:
    url = "https://api.groq.com/openai/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }
    payload = {
        "model": model_choice,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.85,
        "max_tokens": 8192,
    }
    resp = requests.post(url, headers=headers, json=payload, timeout=120)
    if resp.status_code != 200:
        raise Exception(f"Groq API error {resp.status_code}: {resp.text}")
    data = resp.json()
    return data["choices"][0]["message"]["content"], model_choice

# ── OpenRouter ────────────────────────────────────────
OPENROUTER_MODELS = {
    "Meta Llama 3.3 70B (Free)":    "meta-llama/llama-3.3-70b-instruct:free",
    "DeepSeek R1 (Free)":           "deepseek/deepseek-r1:free",
    "Mistral 7B (Free)":            "mistralai/mistral-7b-instruct:free",
    "Gemma 3 27B (Free)":           "google/gemma-3-27b-it:free",
    "Qwen 2.5 72B (Free)":          "qwen/qwen2.5-72b-instruct:free",
}

def generate_openrouter(prompt: str, api_key: str, model_id: str, model_label: str) -> tuple[str, str]:
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://smartquiz-ai.streamlit.app",
        "X-Title": "SmartQuiz AI",
    }
    payload = {
        "model": model_id,
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.85,
        "max_tokens": 8192,
    }
    resp = requests.post(url, headers=headers, json=payload, timeout=120)
    if resp.status_code != 200:
        raise Exception(f"OpenRouter API error {resp.status_code}: {resp.text}")
    data = resp.json()
    return data["choices"][0]["message"]["content"], model_label

# ── Router utama ──────────────────────────────────────
def generate_ai(prompt: str, provider: str, api_key: str, model_pilihan: str) -> tuple[str, str]:
    if provider == "Google Gemini":
        return generate_gemini(prompt, api_key)
    elif provider == "Groq":
        return generate_groq(prompt, api_key, model_pilihan)
    elif provider == "OpenRouter":
        model_id = OPENROUTER_MODELS.get(model_pilihan, list(OPENROUTER_MODELS.values())[0])
        return generate_openrouter(prompt, api_key, model_id, model_pilihan)
    else:
        raise Exception("Provider tidak dikenali.")


# =====================================================
# 5. SIDEBAR
# =====================================================
with st.sidebar:
    st.markdown("""
    <div style="text-align:center; padding:1rem 0 0.5rem;">
        <div style="font-size:2.8rem;">⚡</div>
        <div style="font-size:1.1rem; font-weight:800; color:white; letter-spacing:-0.5px;">SmartQuiz AI</div>
        <div style="font-size:0.7rem; color:#A5B4FC; margin-top:2px; letter-spacing:0.08em;">v3.0 · Multi-Provider AI</div>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("---")

    # ── Pilih Provider ──
    st.markdown("#### 🤖 Pilih Provider AI")
    provider = st.selectbox(
        "Provider",
        ["Google Gemini", "Groq", "OpenRouter"],
        label_visibility="collapsed",
        help="Pilih layanan AI yang ingin digunakan."
    )

    # Info singkat tiap provider
    info_provider = {
        "Google Gemini": "🔵 **Gemini** — Model default. Gratis 1.500 req/hari.\n[→ Dapatkan API Key](https://aistudio.google.com/app/apikey)",
        "Groq":          "🟠 **Groq** — Super cepat! Gratis 14.400 req/hari.\n[→ Dapatkan API Key](https://console.groq.com)",
        "OpenRouter":    "🟢 **OpenRouter** — Banyak model gratis.\n[→ Dapatkan API Key](https://openrouter.ai/keys)",
    }
    st.info(info_provider[provider])

    # ── Pilih Model (khusus Groq & OpenRouter) ──
    model_pilihan = None
    if provider == "Groq":
        st.markdown("#### 🧠 Model Groq")
        model_pilihan = st.selectbox(
            "Model Groq", GROQ_MODELS, label_visibility="collapsed",
            help="Llama 3.3 70B direkomendasikan untuk kualitas terbaik."
        )
    elif provider == "OpenRouter":
        st.markdown("#### 🧠 Model OpenRouter")
        model_pilihan = st.selectbox(
            "Model OpenRouter", list(OPENROUTER_MODELS.keys()), label_visibility="collapsed",
            help="Semua model yang tersedia adalah GRATIS."
        )

    # ── Input API Key ──
    st.markdown("---")
    placeholder_map = {
        "Google Gemini": "AIzaSy...",
        "Groq":          "gsk_...",
        "OpenRouter":    "sk-or-...",
    }
    st.markdown("#### 🔑 API Key")
    api_key = st.text_input(
        f"{provider} API Key",
        type="password",
        placeholder=placeholder_map[provider],
        label_visibility="collapsed"
    )
    if api_key:
        if len(api_key) < 20:
            st.error("⚠️ API Key tampaknya tidak valid.")
        else:
            st.success(f"✅ API Key {provider} terdeteksi")

    # ── Bahasa Output ──
    st.markdown("---")
    st.markdown("#### 🌐 Bahasa Output")
    bahasa_output = st.selectbox(
        "Bahasa", ["Bahasa Indonesia", "Bahasa Inggris", "Bahasa Melayu"],
        label_visibility="collapsed"
    )

    st.markdown("---")
    st.markdown("""
    <div style="font-size:0.82rem; color:#A5B4FC; line-height:1.9;">
        <div style="color:#C7D2FE; font-weight:700; margin-bottom:0.4rem;">✨ Fitur Unggulan</div>
        🤖 &nbsp;3 Provider AI gratis<br>
        🚀 &nbsp;Generate super cepat<br>
        🖼️ &nbsp;Ilustrasi gambar otomatis<br>
        📊 &nbsp;Kisi-kisi & kartu soal<br>
        📄 &nbsp;Export ke Word (.docx)<br>
        📱 &nbsp;Responsif di HP & Tablet
    </div>
    """, unsafe_allow_html=True)


# =====================================================
# 6. FORM INPUT UTAMA
# =====================================================
st.markdown('<div class="section-header">📋 Konfigurasi Evaluasi</div>', unsafe_allow_html=True)

col1, col2 = st.columns(2, gap="large")
with col1:
    mapel = st.text_input("📚 Mata Pelajaran", placeholder="Contoh: Matematika, IPA, Bahasa Indonesia...")
    kelas = st.selectbox("🎓 Kelas", [f"Kelas {i}" for i in range(1, 13)], index=9)
    format_soal = st.selectbox(
        "📝 Format Soal",
        ["Pilihan Ganda", "Pilihan Jamak (>1 Jawaban)", "Benar Salah", "Uraian / Essay", "Menjodohkan"]
    )
    if format_soal != "Uraian / Essay":
        jml_opsi_label = {
            "Benar Salah": "Jumlah Pernyataan",
            "Menjodohkan": "Jumlah Pasangan",
        }.get(format_soal, "Jumlah Opsi Jawaban")
        jml_opsi = st.selectbox(f"🔢 {jml_opsi_label}", [3, 4, 5, 6], index=1)
    else:
        jml_opsi = 0

with col2:
    topik = st.text_area(
        "🎯 Topik / Tujuan Pembelajaran",
        placeholder="Contoh: Peserta didik mampu memahami konsep fotosintesis...",
        height=130
    )
    st.markdown("<div style='margin-top:0.6rem;'></div>", unsafe_allow_html=True)
    mode_bergambar = st.checkbox("🖼️ Sisipkan Ilustrasi Gambar Otomatis", value=False)
    if mode_bergambar:
        st.caption("💡 Gambar diambil otomatis dari internet. Bisa diganti/dihapus di Word.")
    sertakan_rubrik = st.checkbox(
        "📋 Sertakan Rubrik Penilaian (khusus Uraian)", value=False,
        disabled=(format_soal != "Uraian / Essay")
    )

st.markdown("---")

# Komposisi soal
st.markdown('<div class="section-header">📊 Komposisi Tingkat Kesulitan</div>', unsafe_allow_html=True)
c1, c2, c3 = st.columns(3, gap="medium")
with c1: jml_mudah  = st.number_input("🟢 Mudah",  0, 50, 3, 1)
with c2: jml_sedang = st.number_input("🟡 Sedang", 0, 50, 4, 1)
with c3: jml_sulit  = st.number_input("🔴 Sulit",  0, 50, 3, 1)

total_soal = jml_mudah + jml_sedang + jml_sulit

st.markdown("<div style='margin:1.2rem 0 0.5rem;'></div>", unsafe_allow_html=True)
s1, s2, s3, s4 = st.columns(4)
for col, num, label in zip([s1, s2, s3, s4],
                            [total_soal, jml_mudah, jml_sedang, jml_sulit],
                            ["Total Soal", "Mudah", "Sedang", "Sulit"]):
    with col:
        st.markdown(
            f'<div class="stat-card"><div class="stat-number">{num}</div>'
            f'<div class="stat-label">{label}</div></div>',
            unsafe_allow_html=True
        )

st.markdown("<div style='margin-bottom:1.2rem;'></div>", unsafe_allow_html=True)

# Taksonomi Bloom
st.markdown('<div class="section-header">🧠 Level Kognitif (Taksonomi Bloom)</div>', unsafe_allow_html=True)
kognitif = st.multiselect(
    "Pilih level",
    ["C1 – Mengingat", "C2 – Memahami", "C3 – Menerapkan",
     "C4 – Menganalisis", "C5 – Mengevaluasi", "C6 – Mencipta"],
    default=["C2 – Memahami", "C3 – Menerapkan"],
    label_visibility="collapsed"
)

st.markdown("<br>", unsafe_allow_html=True)


# =====================================================
# 7. BUILDER PROMPT
# =====================================================
def build_format_instruksi(format_soal, jml_opsi):
    if format_soal == "Benar Salah":
        return (
            f"Format BENAR/SALAH. Setiap soal berisi {jml_opsi} PERNYATAAN TERPISAH.\n"
            f"- Buat baris baru untuk tiap pernyataan, berlabel a., b., c., dst.\n"
            f"- DILARANG menggabungkan pernyataan dalam satu paragraf.\n"
            f"- Akhiri tiap soal dengan: 'Tentukan apakah setiap pernyataan BENAR atau SALAH.'"
        )
    elif format_soal == "Uraian / Essay":
        return (
            "Format URAIAN. Pertanyaan terbuka, TANPA opsi jawaban.\n"
            "Setiap soal harus mendorong peserta didik berpikir kritis dan terukur."
        )
    elif format_soal == "Pilihan Jamak (>1 Jawaban)":
        return (
            f"Format PILIHAN JAMAK (>1 jawaban bisa benar). Setiap soal: {jml_opsi} opsi.\n"
            f"- Tambahkan petunjuk '(Pilih semua jawaban yang benar)' setelah teks soal.\n"
            f"- Buat baris baru untuk tiap opsi, berlabel A., B., C., dst."
        )
    elif format_soal == "Menjodohkan":
        return (
            f"Format MENJODOHKAN. {jml_opsi} pasangan per soal.\n"
            f"- Kolom A: istilah/pernyataan berlabel 1., 2., 3., dst.\n"
            f"- Kolom B: jawaban berlabel a., b., c., dst. (DIACAK, lebih banyak 1-2 dari Kolom A).\n"
            f"- Instruksikan peserta untuk mencocokkan Kolom A dengan Kolom B."
        )
    else:  # Pilihan Ganda
        return (
            f"Format PILIHAN GANDA (1 jawaban benar). Setiap soal: {jml_opsi} opsi.\n"
            f"- Buat baris baru untuk tiap opsi, berlabel A., B., C., dst.\n"
            f"- Buat pengecoh (distraktor) yang masuk akal."
        )

def build_prompt(mapel, kelas, topik, kognitif, format_soal, jml_opsi,
                 jml_mudah, jml_sedang, jml_sulit, total_soal,
                 mode_bergambar, sertakan_rubrik, bahasa_output):

    instruksi_format = build_format_instruksi(format_soal, jml_opsi)

    instruksi_gambar = (
        "4. ATURAN GAMBAR (WAJIB): Sisipkan {{GAMBAR: kata kunci}} di SETIAP soal.\n"
        "   Kata kunci: 1-2 kata benda bahasa Inggris yang spesifik.\n"
        "   Contoh: {{GAMBAR: human heart}} atau {{GAMBAR: plant cell}}"
        if mode_bergambar
        else "4. JANGAN menyisipkan gambar apapun."
    )

    instruksi_rubrik = (
        "\n\nDi [BAGIAN_KUNCI], sertakan RUBRIK PENILAIAN per soal uraian dengan skor per aspek."
        if sertakan_rubrik and "Uraian" in format_soal
        else ""
    )

    instruksi_bahasa = {
        "Bahasa Indonesia": "Tulis seluruh output dalam Bahasa Indonesia yang baik dan benar.",
        "Bahasa Inggris":   "Write all output in proper English.",
        "Bahasa Melayu":    "Tulis semua output dalam Bahasa Melayu yang betul.",
    }.get(bahasa_output, "")

    kognitif_str = ", ".join(kognitif) if kognitif else "C2 – Memahami, C3 – Menerapkan"

    return f"""
Anda adalah Guru Ahli Kurikulum Merdeka yang berpengalaman.
{instruksi_bahasa}

Buatkan {total_soal} soal evaluasi berkualitas tinggi:
- Mata Pelajaran : {mapel}
- Kelas          : {kelas}
- Topik / TP     : {topik}
- Level Kognitif : {kognitif_str}
- Distribusi     : Mudah ({jml_mudah}), Sedang ({jml_sedang}), Sulit ({jml_sulit})

ATURAN LAYOUT (WAJIB 100%):
1. DILARANG menggunakan Heading Markdown (#, ##, ###). Gunakan **cetak tebal** untuk sub-judul.
2. DILARANG menggunakan format LaTeX atau simbol dolar ($). Tulis rumus teks biasa: x^2, 1/2.
3. {instruksi_format}
{instruksi_gambar}

OUTPUT (ikuti urutan persis ini):

[BAGIAN_SOAL]
(Tulis semua soal evaluasi){instruksi_rubrik}

[BAGIAN_KUNCI]
(Tabel kunci jawaban dan pembahasan singkat.
 Format: | No | Jawaban | Pembahasan Singkat |)

[BAGIAN_KISI]
Tabel kisi-kisi dalam format Markdown:
| No | Indikator Soal | Level Kognitif | Tingkat Kesulitan |
|---|---|---|---|

[BAGIAN_KARTU]
Kartu Soal Kurikulum Merdeka untuk SETIAP soal:
---
**KARTU SOAL NOMOR [X]**
* **Tujuan Pembelajaran:** {topik}
* **Materi Pokok:** [isi]
* **Indikator Soal:** [isi]
* **Level Kognitif:** [isi]
* **Tingkat Kesulitan:** [Mudah/Sedang/Sulit]

**Rumusan Soal:**
[Teks soal lengkap beserta opsi jawaban]
"""


# =====================================================
# 8. TOMBOL GENERATE & LOGIKA UTAMA
# =====================================================

# Inisialisasi session_state agar data tidak hilang saat tombol download diklik
if "hasil_generate" not in st.session_state:
    st.session_state.hasil_generate = None  # Menyimpan semua hasil generate

generate_clicked = st.button("🚀 Generate Evaluasi Sekarang", use_container_width=True, type="primary")

if generate_clicked:
    # ── Validasi ──
    errors = []
    if not api_key:
        errors.append(f"⚠️ Masukkan **{provider} API Key** di sidebar kiri.")
    if not mapel.strip():
        errors.append("⚠️ Kolom **Mata Pelajaran** tidak boleh kosong.")
    if not topik.strip():
        errors.append("⚠️ Kolom **Topik / Tujuan Pembelajaran** tidak boleh kosong.")
    if total_soal == 0:
        errors.append("⚠️ Total soal tidak boleh **0**.")
    if total_soal > 50:
        errors.append("⚠️ Total soal maksimal **50** untuk menjaga kualitas output.")
    if not kognitif:
        errors.append("⚠️ Pilih minimal **1 level kognitif** Taksonomi Bloom.")

    if errors:
        for e in errors:
            st.error(e)
        st.stop()

    prompt = build_prompt(
        mapel, kelas, topik, kognitif, format_soal, jml_opsi,
        jml_mudah, jml_sedang, jml_sulit, total_soal,
        mode_bergambar, sertakan_rubrik, bahasa_output
    )

    try:
        progress_bar = st.progress(0, text=f"🤖 Menghubungi {provider}...")
        time.sleep(0.3)
        progress_bar.progress(10, text=f"📡 Mengirim permintaan ke **{provider}**...")

        with st.spinner(f"📝 {provider} sedang menyusun {total_soal} soal..."):
            hasil, model_used = generate_ai(prompt, provider, api_key, model_pilihan)

        progress_bar.progress(70, text="✨ Soal berhasil di-generate! Memproses hasil...")
        time.sleep(0.3)

        # Bersihkan heading markdown
        hasil_clean = re.sub(r'^#+\s+(.*)$', r'**\1**', hasil, flags=re.MULTILINE)

        # Proses gambar
        if mode_bergambar:
            def ubah_ke_url(match):
                kw = match.group(1).strip()
                kw_aman = urllib.parse.quote(kw.replace(" ", ","))
                return f"\n\n![Ilustrasi {kw}](https://loremflickr.com/600/350/{kw_aman}/all)\n\n"
            hasil_clean = re.sub(r'\{\{GAMBAR:\s*(.*?)\}\}', ubah_ke_url, hasil_clean, flags=re.IGNORECASE)

        # Pisahkan 4 bagian
        parts = re.split(r'\[BAGIAN_SOAL\]|\[BAGIAN_KUNCI\]|\[BAGIAN_KISI\]|\[BAGIAN_KARTU\]', hasil_clean)
        parts = [p for p in parts if p.strip()]

        soal_teks  = parts[0].strip() if len(parts) > 0 else hasil_clean
        kunci_teks = parts[1].strip() if len(parts) > 1 else "*Kunci jawaban tidak berhasil dipisahkan.*"
        kisi_teks  = parts[2].strip() if len(parts) > 2 else "*Kisi-kisi tidak berhasil dipisahkan.*"
        kartu_teks = parts[3].strip() if len(parts) > 3 else "*Kartu soal tidak berhasil dipisahkan.*"

        progress_bar.progress(90, text="📄 Menyiapkan file Word...")
        time.sleep(0.3)

        # Siapkan semua DOCX (dibuat sekali, disimpan di session_state)
        info_doc = {"mapel": mapel, "kelas": kelas, "topik": topik}

        def buat_docx_bagian(judul, konten):
            return export_to_docx(judul, info_doc, konten).getvalue()

        doc_lengkap = buat_docx_bagian(
            f"SOAL EVALUASI LENGKAP: {mapel.upper()} – {kelas}",
            f"{soal_teks}\n\n## Kunci Jawaban & Pembahasan\n{kunci_teks}\n\n## Kisi-Kisi Soal\n{kisi_teks}\n\n## Kumpulan Kartu Soal\n{kartu_teks}"
        )
        doc_soal   = buat_docx_bagian(f"SOAL EVALUASI: {mapel.upper()} – {kelas}", soal_teks)
        doc_kunci  = buat_docx_bagian(f"KUNCI JAWABAN & PEMBAHASAN: {mapel.upper()} – {kelas}", kunci_teks)
        doc_kisi   = buat_docx_bagian(f"KISI-KISI SOAL: {mapel.upper()} – {kelas}", kisi_teks)
        doc_kartu  = buat_docx_bagian(f"KARTU SOAL: {mapel.upper()} – {kelas}", kartu_teks)

        # ✅ Simpan semua data ke session_state agar tidak hilang saat download
        st.session_state.hasil_generate = {
            "soal_teks":   soal_teks,
            "kunci_teks":  kunci_teks,
            "kisi_teks":   kisi_teks,
            "kartu_teks":  kartu_teks,
            "model_used":  model_used,
            "provider":    provider,
            "total_soal":  total_soal,
            "mapel":       mapel,
            "kelas":       kelas,
            "doc_lengkap": doc_lengkap,
            "doc_soal":    doc_soal,
            "doc_kunci":   doc_kunci,
            "doc_kisi":    doc_kisi,
            "doc_kartu":   doc_kartu,
        }

        progress_bar.progress(100, text="🎉 Selesai!")
        time.sleep(0.5)
        progress_bar.empty()

    except Exception as e:
        error_msg = str(e)
        if 'progress_bar' in locals():
            progress_bar.empty()

        if "API_KEY" in error_msg.upper() or "invalid" in error_msg.lower() or "401" in error_msg:
            st.error(f"🔐 **API Key {provider} tidak valid.** Pastikan API Key benar dan aktif.")
        elif "quota" in error_msg.lower() or "429" in error_msg or "rate" in error_msg.lower():
            st.error(f"⏳ **Kuota {provider} habis.** Coba ganti provider di sidebar (Groq / OpenRouter).")
        elif "timeout" in error_msg.lower():
            st.error("⌛ **Koneksi timeout.** Coba kurangi jumlah soal atau ganti provider.")
        elif "insufficient_quota" in error_msg.lower():
            st.error("💳 **Kredit API habis.** Gunakan provider lain yang masih gratis.")
        else:
            st.error(f"❌ **Terjadi kesalahan ({provider}):** {error_msg}")

        st.info("💡 **Solusi cepat:** Ganti provider AI di sidebar kiri — Groq dan OpenRouter tersedia gratis!")
        st.stop()


# =====================================================
# 9. TAMPILKAN HASIL (dari session_state — tidak hilang saat download)
# =====================================================
if st.session_state.hasil_generate:
    r = st.session_state.hasil_generate  # shortcut

    st.success(
        f"✨ Berhasil! **{r['total_soal']} soal** selesai di-generate "
        f"menggunakan **{r['provider']}** · model `{r['model_used']}`."
    )

    st.markdown("<br>", unsafe_allow_html=True)

    # ── Tab Tampilan ──
    tab1, tab2, tab3, tab4 = st.tabs([
        "📄 Soal Evaluasi", "🔑 Kunci & Pembahasan", "📊 Kisi-Kisi", "📇 Kartu Soal"
    ])
    with tab1: st.markdown(r["soal_teks"])
    with tab2: st.markdown(r["kunci_teks"])
    with tab3: st.markdown(r["kisi_teks"])
    with tab4: st.markdown(r["kartu_teks"])

    st.markdown("---")
    st.markdown("### ⬇️ Download File")

    # ── Baris 1: Download Lengkap ──
    dl0_col1, dl0_col2 = st.columns([3, 1])
    with dl0_col1:
        st.download_button(
            label="📦 Download LENGKAP (Soal + Kunci + Kisi + Kartu)",
            data=r["doc_lengkap"],
            file_name=f"Evaluasi_Lengkap_{r['mapel'].replace(' ','_')}_{r['kelas'].replace(' ','_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="dl_lengkap"
        )
    with dl0_col2:
        st.info(f"📁 {len(r['doc_lengkap']) // 1024} KB")

    st.markdown("<div style='margin-top:0.5rem;'></div>", unsafe_allow_html=True)

    # ── Baris 2: Download Terpisah ──
    st.markdown("**Atau download per bagian:**")
    d1, d2, d3, d4 = st.columns(4, gap="small")

    with d1:
        st.download_button(
            label="📄 Soal Saja",
            data=r["doc_soal"],
            file_name=f"Soal_{r['mapel'].replace(' ','_')}_{r['kelas'].replace(' ','_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="dl_soal"
        )
    with d2:
        st.download_button(
            label="🔑 Kunci & Pembahasan",
            data=r["doc_kunci"],
            file_name=f"Kunci_{r['mapel'].replace(' ','_')}_{r['kelas'].replace(' ','_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="dl_kunci"
        )
    with d3:
        st.download_button(
            label="📊 Kisi-Kisi",
            data=r["doc_kisi"],
            file_name=f"KisiKisi_{r['mapel'].replace(' ','_')}_{r['kelas'].replace(' ','_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="dl_kisi"
        )
    with d4:
        st.download_button(
            label="📇 Kartu Soal",
            data=r["doc_kartu"],
            file_name=f"KartuSoal_{r['mapel'].replace(' ','_')}_{r['kelas'].replace(' ','_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
            key="dl_kartu"
        )

    st.markdown("<br>", unsafe_allow_html=True)

    with st.expander("💡 Tips: Merapikan Rumus Matematika di Microsoft Word"):
        st.info("""
**Cara convert rumus teks menjadi Equation rapi di Word:**
1. **Salin** teks rumus dari hasil AI (misal: `x^2 - 4ac = 0`)
2. **Tempel** ke dokumen Microsoft Word
3. **Blok / sorot** teks rumus
4. Tekan **`Alt + =`**
5. Klik **Design** → **Convert** → **Convert to Professional**
        """)

    with st.expander("ℹ️ Tentang Gambar Ilustrasi"):
        st.info("""
Gambar diambil otomatis dari **LoremFlickr** berdasarkan kata kunci AI.
- Di **web**: terlihat langsung di halaman ini.
- Di **Word**: disematkan otomatis saat download.
- Gambar kurang pas? Bisa diganti/dihapus di Microsoft Word.
        """)

    with st.expander("🔄 Panduan Jika Kuota API Habis"):
        st.info("""
| Provider | Limit Gratis | Link Daftar |
|---|---|---|
| 🔵 Google Gemini | 1.500 req/hari | [aistudio.google.com](https://aistudio.google.com/app/apikey) |
| 🟠 Groq | 14.400 req/hari | [console.groq.com](https://console.groq.com) |
| 🟢 OpenRouter | Banyak model gratis | [openrouter.ai](https://openrouter.ai/keys) |

**Solusi cepat:** Ganti provider di sidebar kiri — semua gratis!
        """)
